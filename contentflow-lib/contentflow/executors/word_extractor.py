"""Word document extraction executor using python-docx for text, paragraphs, and tables."""

import base64
import io
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError(
        "python-docx is required for Word document extraction. "
        "Install it with: pip install python-docx"
    )

from . import ParallelExecutor
from ..models import Content

logger = logging.getLogger("contentflow.executors.word_extractor")


class WordExtractorExecutor(ParallelExecutor):
    """
    Extract content from Word documents using python-docx.
    
    This executor analyzes Word documents (.docx) to extract text, paragraphs,
    tables, and document properties using the python-docx library.
    
    Configuration (settings dict):
        - extract_text (bool): Extract full text content from document
          Default: True
        - extract_paragraphs (bool): Create paragraph-level chunks
          Default: True
        - extract_tables (bool): Extract tables from document
          Default: True
        - extract_properties (bool): Extract document properties (author, title, etc.)
          Default: False
        - extract_images (bool): Extract embedded images from document
          Default: False
        - content_field (str): Field containing Word document bytes
          Default: None
        - temp_file_path_field (str): Field containing temp file path
          Default: "temp_file_path"
        - output_field (str): Field name for extracted data
          Default: "word_output"
        - image_output_mode (str): How to store extracted images
          Default: "base64"
          Options: "base64", "bytes"
        - paragraph_separator (str): Separator between paragraphs in full text
          Default: "\n"
        - include_empty_paragraphs (bool): Include empty paragraphs in extraction
          Default: False

        Also setting from ParallelExecutor and BaseExecutor apply.
    
    Example:
        ```python
        executor = WordExtractorExecutor(
            id="word_extractor",
            settings={
                "extract_text": True,
                "extract_paragraphs": True,
                "extract_tables": True,
                "extract_properties": True,
            }
        )
        ```
    
    Input:
        Document or List[Document] each with (ContentIdentifier) id containing:
        - data['content']: Word document bytes, OR
        - data['temp_file_path']: Path to Word document file
        
    Output:
        Document or List[Document] with added fields:
        - data['word_output']['text']: Full extracted text
        - data['word_output']['paragraphs']: List of paragraph chunks with text and metadata
        - data['word_output']['tables']: List of extracted tables (if enabled)
        - data['word_output']['properties']: Document properties (if enabled)
        - data['word_output']['images']: List of extracted images (if enabled)
        - summary_data['paragraphs_processed']: Number of paragraphs processed
        - summary_data['tables_extracted']: Number of tables extracted
    """
    
    def __init__(
        self,
        id: str,
        settings: Optional[Dict[str, Any]] = None,
        **kwargs
    ):
        super().__init__(
            id=id,
            settings=settings,
            **kwargs
        )
        
        # Extract configuration
        self.extract_text = self.get_setting("extract_text", default=True)
        self.extract_paragraphs = self.get_setting("extract_paragraphs", default=True)
        self.extract_tables = self.get_setting("extract_tables", default=True)
        self.extract_properties = self.get_setting("extract_properties", default=False)
        self.extract_images = self.get_setting("extract_images", default=False)
        self.content_field = self.get_setting("content_field", default=None)
        self.temp_file_field = self.get_setting("temp_file_path_field", default="temp_file_path")
        self.output_field = self.get_setting("output_field", default="word_output")
        self.image_output_mode = self.get_setting("image_output_mode", default="base64")
        self.paragraph_separator = self.get_setting("paragraph_separator", default="\n")
        self.include_empty_paragraphs = self.get_setting("include_empty_paragraphs", default=False)
        
        # Validate image output mode
        if self.image_output_mode not in ["base64", "bytes"]:
            raise ValueError(f"Invalid image_output_mode: {self.image_output_mode}. Must be 'base64' or 'bytes'")
        
        if self.debug_mode:
            logger.debug(
                f"WordExtractorExecutor with id {self.id} initialized: "
                f"text={self.extract_text}, paragraphs={self.extract_paragraphs}, "
                f"tables={self.extract_tables}, properties={self.extract_properties}"
            )
    
    async def process_content_item(
        self,
        content: Content
    ) -> Content:
        """Process a single Word document using python-docx.
           Implements the abstract method from ParallelExecutor.
        """
        
        try:
            if not content or not content.data:
                raise ValueError("Content must have data")
            
            # Get Word document content
            doc_bytes = None
            doc_path = None
            
            # Try to get from content field
            if self.content_field and self.content_field in content.data:
                doc_bytes = content.data[self.content_field]
            
            # Try to get from temp file
            elif self.temp_file_field in content.data:
                doc_path = content.data[self.temp_file_field]
            
            if not doc_bytes and not doc_path:
                raise ValueError(
                    f"Word document missing required content. "
                    f"Needs either '{self.content_field}' or '{self.temp_file_field}'"
                )
            
            if self.debug_mode:
                source = f"file: {doc_path}" if doc_path else f"bytes: {len(doc_bytes)} bytes"
                logger.debug(f"Processing Word document {content.id} from {source}")
            
            doc = None
            
            try:
                # Open Word document
                if doc_bytes:
                    doc = Document(io.BytesIO(doc_bytes))
                else:
                    doc = Document(doc_path)
            except Exception as e:
                logger.warning(
                    f"Invalid Word document for content {content.id}: {str(e)}"
                )
                content.summary_data['word_extraction_status'] = "invalid_file"
                return content
            
            extracted_data = {}
            
            # Extract paragraphs
            all_text = []
            paragraphs_data = []
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                
                # Skip empty paragraphs if configured
                if not self.include_empty_paragraphs and not para_text.strip():
                    continue
                
                if self.extract_text:
                    all_text.append(para_text)
                
                if self.extract_paragraphs:
                    para_info = {
                        "paragraph_number": para_num + 1,
                        "text": para_text,
                        "char_count": len(para_text),
                        "style": paragraph.style.name if paragraph.style else None,
                    }
                    paragraphs_data.append(para_info)
            
            if self.extract_text:
                extracted_data['text'] = self.paragraph_separator.join(all_text)
                if self.debug_mode:
                    logger.debug(f"Extracted {len(extracted_data['text'])} characters of text")
            
            if self.extract_paragraphs:
                extracted_data['paragraphs'] = paragraphs_data
                if self.debug_mode:
                    logger.debug(f"Created {len(paragraphs_data)} paragraph chunks")
            
            # Extract tables
            tables_count = 0
            if self.extract_tables:
                tables = self._extract_tables_from_document(doc)
                extracted_data['tables'] = tables
                tables_count = len(tables)
                if self.debug_mode:
                    logger.debug(f"Extracted {tables_count} tables")
            
            # Extract document properties
            if self.extract_properties:
                properties = self._extract_document_properties(doc)
                extracted_data['properties'] = properties
                if self.debug_mode:
                    logger.debug(f"Extracted document properties")
            
            # Extract images
            images_count = 0
            if self.extract_images:
                images = self._extract_images_from_document(doc)
                extracted_data['images'] = images
                images_count = len(images)
                if self.debug_mode:
                    logger.debug(f"Extracted {images_count} images")
            
            # Store extracted data
            content.data[self.output_field] = extracted_data
            
            # Update summary
            content.summary_data['paragraphs_processed'] = len(paragraphs_data) if self.extract_paragraphs else 0
            content.summary_data['tables_extracted'] = tables_count
            content.summary_data['images_extracted'] = images_count
            content.summary_data['extraction_status'] = "success"
                            
        except Exception as e:
            logger.error(
                f"WordExtractorExecutor {self.id} failed processing document {content.id}",
                exc_info=True
            )
            
            # Raise the exception to be handled upstream if needed
            raise
        
        return content
    
    def _extract_tables_from_document(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract tables from Word document.
        
        Args:
            doc: Opened python-docx Document
            
        Returns:
            List of table dictionaries with cell data
        """
        tables = []
        
        for table_num, table in enumerate(doc.tables):
            try:
                table_data = {
                    "table_number": table_num + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": []
                }
                
                # Extract cell data
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data["data"].append(row_data)
                
                tables.append(table_data)
                
            except Exception as e:
                logger.warning(f"Failed to extract table {table_num + 1}: {e}")
                continue
        
        return tables
    
    def _extract_document_properties(self, doc: Document) -> Dict[str, Any]:
        """Extract document properties and metadata.
        
        Args:
            doc: Opened python-docx Document
            
        Returns:
            Dictionary of document properties
        """
        properties = {}
        
        try:
            core_props = doc.core_properties
            
            properties['title'] = core_props.title or None
            properties['author'] = core_props.author or None
            properties['subject'] = core_props.subject or None
            properties['keywords'] = core_props.keywords or None
            properties['comments'] = core_props.comments or None
            properties['category'] = core_props.category or None
            properties['created'] = core_props.created.isoformat() if core_props.created else None
            properties['modified'] = core_props.modified.isoformat() if core_props.modified else None
            properties['last_modified_by'] = core_props.last_modified_by or None
            properties['revision'] = core_props.revision
            
        except Exception as e:
            logger.warning(f"Failed to extract some document properties: {e}")
        
        return properties
    
    def _extract_images_from_document(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract images from Word document.
        
        Args:
            doc: Opened python-docx Document
            
        Returns:
            List of image dictionaries with metadata and image data
        """
        images = []
        
        try:
            # Access document parts to find images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        
                        # Determine image format from content type
                        content_type = image_part.content_type
                        image_format = content_type.split('/')[-1] if '/' in content_type else 'unknown'
                        
                        # Prepare image data based on output mode
                        if self.image_output_mode == "base64":
                            image_data = base64.b64encode(image_bytes).decode('utf-8')
                        else:
                            image_data = image_bytes
                        
                        image_info = {
                            "image_index": len(images),
                            "format": image_format,
                            "size_bytes": len(image_bytes),
                            "content_type": content_type,
                            "data": image_data,
                            "encoding": self.image_output_mode
                        }
                        
                        images.append(image_info)
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")
                        continue
                        
        except Exception as e:
            logger.warning(f"Failed to access document images: {e}")
        
        return images
